"""
Advanced PDF OCR System - Portfolio Project
============================================
A professional-grade PDF text extraction system with batch processing,
intelligent format detection, and comprehensive analytics.

Author: Your Name
Tech Stack: Python, Streamlit, Tesseract OCR, SQLite, OpenPyXL
Features: Single/Batch Processing, Format Detection, History, Analytics
"""

import streamlit as st
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image, ImageEnhance, ImageFilter
import io
import zipfile
import json
import re
from datetime import datetime
import sqlite3
import pandas as pd
from docx import Document
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

# ============================================================================
# PAGE CONFIGURATION
# ============================================================================
st.set_page_config(
    page_title="PDF OCR System Pro",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================================
# CUSTOM STYLING
# ============================================================================
st.markdown("""
    <style>
    .main-header {
        font-size: 2.8rem;
        font-weight: bold;
        background: linear-gradient(120deg, #1f77b4, #ff7f0e, #2ca02c);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #666;
        text-align: center;
        margin-bottom: 2rem;
    }
    .feature-badge {
        display: inline-block;
        padding: 0.25rem 0.7rem;
        margin: 0.2rem;
        border-radius: 12px;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        font-size: 0.8rem;
        font-weight: 600;
    }
    .success-box {
        background: linear-gradient(90deg, #11998e 0%, #38ef7d 100%);
        color: white;
        padding: 1.2rem;
        border-radius: 8px;
        text-align: center;
        font-weight: bold;
        margin: 1rem 0;
    }
    .info-box {
        background: #e3f2fd;
        border-left: 4px solid #2196f3;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    .error-box {
        background: #ffebee;
        border-left: 4px solid #f44336;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

# ============================================================================
# DATABASE FUNCTIONS
# ============================================================================
@st.cache_resource
def init_database():
    """Initialize SQLite database for extraction history"""
    try:
        conn = sqlite3.connect('ocr_history.db', check_same_thread=False)
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS extractions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT NOT NULL,
                timestamp TEXT NOT NULL,
                page_count INTEGER,
                word_count INTEGER,
                character_count INTEGER,
                language TEXT,
                extracted_text TEXT,
                processing_time REAL,
                quality_score REAL
            )
        ''')
        conn.commit()
        return conn
    except Exception as e:
        st.error(f"Database error: {e}")
        return None

def save_to_database(conn, filename, text, metadata):
    """Save extraction results to database"""
    try:
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO extractions 
            (filename, timestamp, page_count, word_count, character_count, 
             language, extracted_text, processing_time, quality_score)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            filename,
            datetime.now().isoformat(),
            metadata.get('page_count', 0),
            metadata.get('word_count', 0),
            metadata.get('character_count', 0),
            metadata.get('language', 'eng'),
            text,
            metadata.get('processing_time', 0),
            metadata.get('quality_score', 0)
        ))
        conn.commit()
        return True
    except Exception as e:
        st.warning(f"Database save failed: {e}")
        return False

# ============================================================================
# IMAGE PREPROCESSING
# ============================================================================
def preprocess_image(image, enhance=True, denoise=True):
    """Apply image enhancement for better OCR accuracy"""
    try:
        if enhance:
            # Convert to grayscale
            image = image.convert('L')
            
            # Enhance contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            
            # Enhance sharpness
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(1.5)
            
            # Reduce noise
            if denoise:
                image = image.filter(ImageFilter.MedianFilter(size=3))
        
        return image
    except Exception as e:
        return image

# ============================================================================
# QUALITY ASSESSMENT
# ============================================================================
def calculate_quality_score(text):
    """Calculate extraction quality based on text characteristics"""
    score = 100.0
    
    # Check text length
    if len(text.strip()) < 10:
        score -= 50
    
    # Check alphanumeric ratio
    alphanumeric_ratio = sum(c.isalnum() or c.isspace() for c in text) / max(len(text), 1)
    if alphanumeric_ratio < 0.7:
        score -= 30
    
    # Check word formation
    words = text.split()
    if words:
        avg_word_length = sum(len(w) for w in words) / len(words)
        if avg_word_length < 2 or avg_word_length > 15:
            score -= 20
    
    return max(0.0, min(100.0, score))

# ============================================================================
# FORMAT DETECTION
# ============================================================================
def detect_document_format(text):
    """Detect document type (invoice, table, form)"""
    formats_detected = []
    
    # Detect tables (multiple columns with numbers)
    lines = text.split('\n')
    table_lines = sum(1 for line in lines if len(line.split()) >= 3 and any(c.isdigit() for c in line))
    if table_lines >= 3:
        formats_detected.append(f"📊 Table ({table_lines} rows)")
    
    # Detect invoice patterns
    invoice_patterns = [
        r'invoice\s*#?\s*:?\s*(\w+)',
        r'total\s*:?\s*\$',
        r'amount\s*due'
    ]
    if any(re.search(pattern, text, re.IGNORECASE) for pattern in invoice_patterns):
        formats_detected.append("💰 Invoice")
    
    # Detect form fields
    form_pattern = r'([A-Z][a-zA-Z\s]+):\s*[_\s]{2,}'
    form_fields = re.findall(form_pattern, text)
    if len(form_fields) >= 3:
        formats_detected.append(f"📋 Form ({len(form_fields)} fields)")
    
    return formats_detected

def extract_key_data(text):
    """Extract key information (emails, phones, dates, amounts)"""
    data = {}
    
    # Extract emails
    emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    if emails:
        data['Emails'] = emails[:3]
    
    # Extract phone numbers
    phones = re.findall(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', text)
    if phones:
        data['Phones'] = phones[:3]
    
    # Extract dates
    dates = re.findall(r'\d{1,2}[-/]\d{1,2}[-/]\d{2,4}', text)
    if dates:
        data['Dates'] = dates[:3]
    
    # Extract amounts
    amounts = re.findall(r'\$\s*(\d{1,3}(?:,\d{3})*(?:\.\d{2})?)', text)
    if amounts:
        data['Amounts'] = amounts[:5]
    
    return data

# ============================================================================
# SEARCH FUNCTIONALITY
# ============================================================================
def search_in_text(text, query, case_sensitive=False):
    """Search for text patterns and return matches with line numbers"""
    if not query:
        return []
    
    matches = []
    lines = text.split('\n')
    
    for i, line in enumerate(lines, 1):
        search_line = line if case_sensitive else line.lower()
        search_query = query if case_sensitive else query.lower()
        
        if search_query in search_line:
            matches.append((i, line))
    
    return matches

# ============================================================================
# EXPORT FUNCTIONS
# ============================================================================
def export_to_word(text, filename, metadata=None):
    """Export to Word document with metadata"""
    try:
        doc = Document()
        
        # Title
        doc.add_heading('OCR Extraction Results', 0)
        
        # Metadata
        doc.add_paragraph(f'Source: {filename}')
        doc.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        if metadata:
            doc.add_paragraph(f'Pages: {metadata.get("page_count", "N/A")}')
            doc.add_paragraph(f'Words: {metadata.get("word_count", "N/A")}')
            doc.add_paragraph(f'Quality: {metadata.get("quality_score", "N/A"):.1f}%')
        
        doc.add_paragraph('')
        doc.add_heading('Extracted Text', level=1)
        
        # Content
        for para in text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
    except Exception as e:
        st.error(f"Word export failed: {e}")
        return None

def export_to_json(text, filename, metadata=None):
    """Export to JSON with metadata"""
    try:
        data = {
            'filename': filename,
            'extracted_text': text,
            'extraction_date': datetime.now().isoformat(),
            'metadata': metadata or {},
            'version': '1.0'
        }
        return json.dumps(data, indent=2)
    except Exception as e:
        st.error(f"JSON export failed: {e}")
        return None

def export_to_excel(text, filename):
    """Export to Excel with analysis sheet"""
    try:
        wb = openpyxl.Workbook()
        
        # Text sheet
        ws1 = wb.active
        ws1.title = "Extracted Text"
        ws1['A1'] = 'Line'
        ws1['B1'] = 'Content'
        ws1['A1'].font = Font(bold=True, color="FFFFFF")
        ws1['B1'].font = Font(bold=True, color="FFFFFF")
        ws1['A1'].fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        ws1['B1'].fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        
        lines = text.split('\n')
        for i, line in enumerate(lines, start=2):
            ws1[f'A{i}'] = i - 1
            ws1[f'B{i}'] = line
        
        ws1.column_dimensions['A'].width = 10
        ws1.column_dimensions['B'].width = 100
        
        # Analysis sheet
        ws2 = wb.create_sheet("Statistics")
        ws2['A1'] = 'Metric'
        ws2['B1'] = 'Value'
        ws2['A1'].font = Font(bold=True)
        ws2['B1'].font = Font(bold=True)
        
        stats = [
            ('Total Lines', len(lines)),
            ('Total Words', len(text.split())),
            ('Total Characters', len(text)),
            ('Non-empty Lines', len([l for l in lines if l.strip()])),
            ('Average Line Length', len(text) // max(len(lines), 1))
        ]
        
        for i, (metric, value) in enumerate(stats, start=2):
            ws2[f'A{i}'] = metric
            ws2[f'B{i}'] = value
        
        ws2.column_dimensions['A'].width = 25
        ws2.column_dimensions['B'].width = 15
        
        # Save
        excel_io = io.BytesIO()
        wb.save(excel_io)
        excel_io.seek(0)
        return excel_io
    except Exception as e:
        st.error(f"Excel export failed: {e}")
        return None

# ============================================================================
# CORE OCR FUNCTION
# ============================================================================
def extract_text_from_pdf(pdf_bytes, ocr_lang, dpi, psm_mode, enhance, denoise):
    """Main OCR extraction function"""
    try:
        # Convert PDF to images
        images = convert_from_bytes(pdf_bytes, dpi=dpi, fmt='jpeg')
        
        extracted_text = ""
        custom_config = f'--oem 3 --psm {psm_mode}'
        
        # Process each page
        for idx, image in enumerate(images):
            # Preprocess
            if enhance:
                image = preprocess_image(image, enhance=True, denoise=denoise)
            
            # Extract text
            try:
                page_text = pytesseract.image_to_string(
                    image,
                    lang=ocr_lang,
                    config=custom_config
                )
                extracted_text += f"\n{'='*60}\nPAGE {idx + 1}\n{'='*60}\n\n"
                extracted_text += page_text.strip() + "\n\n"
            except Exception as e:
                extracted_text += f"\n{'='*60}\nPAGE {idx + 1} - ERROR\n{'='*60}\n\n"
        
        return extracted_text, len(images), None
    except Exception as e:
        return "", 0, str(e)

# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================
if 'current_text' not in st.session_state:
    st.session_state.current_text = ""
if 'processing_stats' not in st.session_state:
    st.session_state.processing_stats = {}

# Initialize database
db_conn = init_database()

# ============================================================================
# HEADER
# ============================================================================
st.markdown('<div class="main-header">🚀 PDF OCR System Pro</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Enterprise-Grade Text Extraction with AI-Powered Analytics</div>', unsafe_allow_html=True)

# Feature badges
st.markdown("""
<div style="text-align: center; margin-bottom: 2rem;">
    <span class="feature-badge">📄 Single Processing</span>
    <span class="feature-badge">📚 Batch Processing</span>
    <span class="feature-badge">🎯 Format Detection</span>
    <span class="feature-badge">🔍 Smart Search</span>
    <span class="feature-badge">💾 History</span>
    <span class="feature-badge">📊 Analytics</span>
</div>
""", unsafe_allow_html=True)

# ============================================================================
# SIDEBAR CONFIGURATION
# ============================================================================
with st.sidebar:
    st.header("⚙️ Configuration")
    
    # Mode selection
    mode = st.radio(
        "📍 Mode",
        ["🔍 Single PDF", "📚 Batch Processing", "📜 History", "📊 Analytics"],
        help="Choose operation mode"
    )
    
    st.divider()
    
    # OCR Settings
    st.subheader("🎛️ OCR Settings")
    
    ocr_lang = st.selectbox(
        "Language",
        ["eng", "fra", "deu", "spa", "ita", "por", "chi_sim", "jpn"],
        index=0
    )
    
    dpi = st.slider("Image Quality (DPI)", 100, 400, 300, 50)
    enhance_image = st.checkbox("🎨 Image Enhancement", value=True)
    denoise = st.checkbox("🧹 Noise Reduction", value=True)
    
    psm_mode = st.selectbox(
        "Page Segmentation",
        [(3, "Fully automatic"), (6, "Uniform block"), (11, "Sparse text")],
        format_func=lambda x: x[1],
        index=0
    )[0]
    
    st.divider()
    
    # Advanced
    st.subheader("🎯 Advanced")
    detect_formats = st.checkbox("📋 Format Detection", value=True)
    save_to_db = st.checkbox("💾 Save to Database", value=True)
    
    st.divider()
    
    # Stats
    if db_conn:
        try:
            cursor = db_conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM extractions")
            total = cursor.fetchone()[0]
            st.metric("Total Extractions", total)
            
            cursor.execute("SELECT SUM(page_count) FROM extractions")
            pages = cursor.fetchone()[0] or 0
            st.metric("Pages Processed", pages)
        except:
            pass

# ============================================================================
# SINGLE PDF MODE
# ============================================================================
if mode == "🔍 Single PDF":
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📤 Upload & Extract")
        
        uploaded_file = st.file_uploader("Choose PDF", type=['pdf'])
        
        if uploaded_file:
            st.success(f"✅ {uploaded_file.name} ({uploaded_file.size/1024:.1f} KB)")
            
            if st.button("🚀 Extract Text", type="primary", use_container_width=True):
                start_time = datetime.now()
                
                with st.spinner("🔄 Processing..."):
                    pdf_bytes = uploaded_file.read()
                    
                    # Extract text
                    progress = st.progress(0)
                    status = st.empty()
                    
                    status.text("Converting PDF to images...")
                    progress.progress(0.3)
                    
                    extracted_text, page_count, error = extract_text_from_pdf(
                        pdf_bytes, ocr_lang, dpi, psm_mode, enhance_image, denoise
                    )
                    
                    progress.progress(1.0)
                    status.empty()
                    progress.empty()
                    
                    if error:
                        st.markdown(f'<div class="error-box">❌ Error: {error}</div>', unsafe_allow_html=True)
                    else:
                        processing_time = (datetime.now() - start_time).total_seconds()
                        
                        # Store results
                        st.session_state.current_text = extracted_text
                        
                        # Calculate metadata
                        metadata = {
                            'filename': uploaded_file.name,
                            'page_count': page_count,
                            'word_count': len(extracted_text.split()),
                            'character_count': len(extracted_text),
                            'language': ocr_lang,
                            'processing_time': processing_time,
                            'quality_score': calculate_quality_score(extracted_text)
                        }
                        
                        st.session_state.processing_stats = metadata
                        
                        # Save to database
                        if save_to_db and db_conn:
                            save_to_database(db_conn, uploaded_file.name, extracted_text, metadata)
                        
                        st.markdown('<div class="success-box">✅ Extraction Complete!</div>', unsafe_allow_html=True)
                        st.balloons()
    
    with col2:
        st.header("📝 Results")
        
        if st.session_state.current_text:
            stats = st.session_state.processing_stats
            
            # Metrics
            col_m1, col_m2, col_m3, col_m4 = st.columns(4)
            col_m1.metric("📄 Pages", stats.get('page_count', 0))
            col_m2.metric("📝 Words", f"{stats.get('word_count', 0):,}")
            col_m3.metric("⏱️ Time", f"{stats.get('processing_time', 0):.1f}s")
            col_m4.metric("✅ Quality", f"{stats.get('quality_score', 0):.0f}%")
            
            # Format detection
            if detect_formats:
                formats = detect_document_format(st.session_state.current_text)
                if formats:
                    st.info("**Detected:** " + " • ".join(formats))
                
                key_data = extract_key_data(st.session_state.current_text)
                if key_data:
                    with st.expander("🔍 Extracted Key Data"):
                        for key, values in key_data.items():
                            st.text(f"{key}: {', '.join(map(str, values))}")
            
            # Text preview
            st.subheader("📄 Extracted Text")
            edited_text = st.text_area(
                "Preview (editable)",
                value=st.session_state.current_text,
                height=200
            )
            
            # Search
            st.subheader("🔍 Search")
            col_s1, col_s2 = st.columns([3, 1])
            with col_s1:
                search_query = st.text_input("Search term", placeholder="Enter text...")
            with col_s2:
                case_sensitive = st.checkbox("Case")
            
            if search_query:
                matches = search_in_text(edited_text, search_query, case_sensitive)
                if matches:
                    st.success(f"✅ Found {len(matches)} matches")
                    with st.expander("View matches"):
                        for line_num, line in matches[:10]:
                            st.text(f"Line {line_num}: {line[:100]}...")
                else:
                    st.warning("No matches found")
            
            # Export
            st.subheader("📥 Export")
            col_e1, col_e2, col_e3 = st.columns(3)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            with col_e1:
                st.download_button(
                    "📄 TXT",
                    data=edited_text,
                    file_name=f"extract_{ts}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
            
            with col_e2:
                word_doc = export_to_word(edited_text, uploaded_file.name, stats)
                if word_doc:
                    st.download_button(
                        "📘 DOCX",
                        data=word_doc,
                        file_name=f"extract_{ts}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            
            with col_e3:
                json_data = export_to_json(edited_text, uploaded_file.name, stats)
                if json_data:
                    st.download_button(
                        "📊 JSON",
                        data=json_data,
                        file_name=f"extract_{ts}.json",
                        mime="application/json",
                        use_container_width=True
                    )
        else:
            st.markdown('<div class="info-box">👆 Upload a PDF and click Extract to begin</div>', unsafe_allow_html=True)

# ============================================================================
# BATCH PROCESSING MODE
# ============================================================================
elif mode == "📚 Batch Processing":
    st.header("📚 Batch Processing")
    st.info("💡 Process multiple PDFs simultaneously")
    
    uploaded_files = st.file_uploader(
        "Choose PDF files",
        type=['pdf'],
        accept_multiple_files=True
    )
    
    if uploaded_files:
        st.success(f"✅ {len(uploaded_files)} files selected")
        
        with st.expander("📋 File List"):
            for i, f in enumerate(uploaded_files, 1):
                st.text(f"{i}. {f.name} ({f.size/1024:.1f} KB)")
        
        if st.button("🚀 Process All", type="primary", use_container_width=True):
            start_time = datetime.now()
            results = {}
            errors = {}
            
            progress = st.progress(0)
            status = st.empty()
            
            for idx, file in enumerate(uploaded_files):
                status.info(f"Processing {idx+1}/{len(uploaded_files)}: {file.name}")
                
                try:
                    pdf_bytes = file.read()
                    text, pages, error = extract_text_from_pdf(
                        pdf_bytes, ocr_lang, dpi, psm_mode, enhance_image, denoise
                    )
                    
                    if error:
                        errors[file.name] = error
                    else:
                        results[file.name] = {
                            'text': text,
                            'pages': pages,
                            'words': len(text.split())
                        }
                        
                        if save_to_db and db_conn:
                            metadata = {
                                'page_count': pages,
                                'word_count': len(text.split()),
                                'character_count': len(text),
                                'language': ocr_lang,
                                'processing_time': 0,
                                'quality_score': calculate_quality_score(text)
                            }
                            save_to_database(db_conn, file.name, text, metadata)
                
                except Exception as e:
                    errors[file.name] = str(e)
                
                progress.progress((idx + 1) / len(uploaded_files))
            
            total_time = (datetime.now() - start_time).total_seconds()
            status.empty()
            progress.empty()
            
            st.markdown('<div class="success-box">✅ Batch Processing Complete!</div>', unsafe_allow_html=True)
            
            # Summary
            col1, col2, col3, col4 = st.columns(4)
            col1.metric("📄 Total", len(uploaded_files))
            col2.metric("✅ Success", len(results))
            col3.metric("❌ Errors", len(errors))
            col4.metric("⏱️ Time", f"{total_time:.1f}s")
            
            if errors:
                with st.expander("⚠️ Errors"):
                    for name, err in errors.items():
                        st.error(f"{name}: {err}")
            
            # Results
            st.subheader("📝 Results")
            for name, data in results.items():
                with st.expander(f"📄 {name}"):
                    col_a, col_b = st.columns(2)
                    col_a.metric("Pages", data['pages'])
                    col_b.metric("Words", data['words'])
                    st.text_area("Preview", data['text'][:500] + "...", height=100, key=f"prev_{name}")
            
            # Export ZIP
            st.subheader("📥 Download All")
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for name, data in results.items():
                    base = name.rsplit('.', 1)[0]
                    zip_file.writestr(f"{base}.txt", data['text'])
            
            zip_buffer.seek(0)
            st.download_button(
                "📦 Download ZIP",
                data=zip_buffer,
                file_name=f"batch_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                mime="application/zip",
                use_container_width=True
            )

# ============================================================================
# HISTORY MODE
# ============================================================================
elif mode == "📜 History":
    st.header("📜 Extraction History")
    
    if db_conn:
        try:
            cursor = db_conn.cursor()
            cursor.execute("""
                SELECT id, filename, timestamp, page_count, word_count, 
                       character_count, language, quality_score 
                FROM extractions 
                ORDER BY timestamp DESC
            """)
            records = cursor.fetchall()
            
            if records:
                # Summary
                col1, col2, col3 = st.columns(3)
                col1.metric("📚 Extractions", len(records))
                col2.metric("📄 Pages", sum(r[3] or 0 for r in records))
                col3.metric("📝 Words", f"{sum(r[4] or 0 for r in records):,}")
                
                st.divider()
                
                # Table
                df = pd.DataFrame(records, columns=[
                    'ID', 'Filename', 'Timestamp', 'Pages', 'Words',
                    'Characters', 'Language', 'Quality'
                ])
                df['Timestamp'] = pd.to_datetime(df['Timestamp']).dt.strftime('%Y-%m-%d %H:%M')
                
                st.dataframe(df, use_container_width=True, hide_index=True)
                
                # Export
                col_exp1, col_exp2 = st.columns(2)
                with col_exp1:
                    csv = df.to_csv(index=False)
                    st.download_button(
                        "📥 Export CSV",
                        data=csv,
                        file_name=f"history_{datetime.now().strftime('%Y%m%d')}.csv",
                        mime="text/csv",
                        use_container_width=True
                    )
                
                with col_exp2:
                    excel_buffer = io.BytesIO()
                    df.to_excel(excel_buffer, index=False, engine='openpyxl')
                    excel_buffer.seek(0)
                    st.download_button(
                        "📗 Export Excel",
                        data=excel_buffer,
                        file_name=f"history_{datetime.now().strftime('%Y%m%d')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                
                st.divider()
                
                # View individual record
                st.subheader("🔍 View Details")
                selected_id = st.selectbox(
                    "Select extraction",
                    options=df['ID'].tolist(),
                    format_func=lambda x: f"#{x} - {df[df['ID']==x]['Filename'].values[0]}"
                )
                
                if selected_id:
                    cursor.execute("SELECT * FROM extractions WHERE id = ?", (selected_id,))
                    record = cursor.fetchone()
                    
                    if record:
                        col_r1, col_r2 = st.columns(2)
                        with col_r1:
                            st.markdown(f"**File:** {record[1]}")
                            st.markdown(f"**Date:** {record[2]}")
                            st.markdown(f"**Pages:** {record[3]}")
                        with col_r2:
                            st.markdown(f"**Words:** {record[4]}")
                            st.markdown(f"**Language:** {record[6]}")
                            st.markdown(f"**Quality:** {record[9]:.1f}%" if record[9] else "N/A")
                        
                        st.text_area("Extracted Text", record[7], height=250, key=f"view_{selected_id}")
                        
                        col_d1, col_d2 = st.columns(2)
                        with col_d1:
                            st.download_button(
                                "📄 Download Text",
                                data=record[7],
                                file_name=f"{record[1]}.txt",
                                use_container_width=True
                            )
                        with col_d2:
                            if st.button("🗑️ Delete", use_container_width=True):
                                cursor.execute("DELETE FROM extractions WHERE id = ?", (selected_id,))
                                db_conn.commit()
                                st.success("✅ Deleted!")
                                st.rerun()
                
                st.divider()
                
                # Clear all
                st.subheader("⚠️ Danger Zone")
                if st.button("🗑️ Clear All History"):
                    if st.checkbox("Confirm deletion of all records"):
                        cursor.execute("DELETE FROM extractions")
                        db_conn.commit()
                        st.success("✅ History cleared!")
                        st.rerun()
            else:
                st.info("📭 No history yet. Process some PDFs to see them here!")
        
        except Exception as e:
            st.error(f"Error: {e}")
    else:
        st.error("❌ Database unavailable")

# ============================================================================
# ANALYTICS MODE
# ============================================================================
elif mode == "📊 Analytics":
    st.header("📊 Analytics Dashboard")
    
    if db_conn:
        try:
            cursor = db_conn.cursor()
            cursor.execute("SELECT * FROM extractions ORDER BY timestamp DESC")
            records = cursor.fetchall()
            
            if records:
                df = pd.DataFrame(records, columns=[
                    'id', 'filename', 'timestamp', 'page_count', 'word_count',
                    'character_count', 'language', 'extracted_text',
                    'processing_time', 'quality_score'
                ])
                
                df['timestamp'] = pd.to_datetime(df['timestamp'])
                df['date'] = df['timestamp'].dt.date
                
                # Overview
                st.subheader("📈 Overview")
                col1, col2, col3, col4 = st.columns(4)
                col1.metric("📚 Total Extractions", len(df))
                col2.metric("📄 Total Pages", int(df['page_count'].sum()))
                col3.metric("📝 Total Words", f"{int(df['word_count'].sum()):,}")
                
                avg_quality = df['quality_score'].mean()
                col4.metric("✅ Avg Quality", f"{avg_quality:.1f}%" if avg_quality else "N/A")
                
                st.divider()
                
                # Extractions over time
                st.subheader("📅 Activity Over Time")
                daily = df.groupby('date').size().reset_index(name='count')
                st.bar_chart(daily.set_index('date'))
                
                st.divider()
                
                # Language distribution
                st.subheader("🌍 Languages Used")
                lang_counts = df['language'].value_counts()
                
                col_lang1, col_lang2 = st.columns([2, 1])
                with col_lang1:
                    st.bar_chart(lang_counts)
                with col_lang2:
                    st.markdown("**Distribution:**")
                    for lang, count in lang_counts.items():
                        st.text(f"{lang}: {count} ({count/len(df)*100:.1f}%)")
                
                st.divider()
                
                # Page distribution
                st.subheader("📄 Pages Per Document")
                page_bins = pd.cut(df['page_count'], bins=[0, 5, 10, 20, 50, 100, 1000])
                page_dist = page_bins.value_counts().sort_index()
                st.bar_chart(page_dist)
                
                st.divider()
                
                # Quality distribution
                if df['quality_score'].notna().any():
                    st.subheader("✅ Quality Score Distribution")
                    quality_bins = pd.cut(
                        df['quality_score'].dropna(),
                        bins=[0, 50, 70, 85, 100],
                        labels=['Poor (0-50)', 'Fair (50-70)', 'Good (70-85)', 'Excellent (85-100)']
                    )
                    quality_dist = quality_bins.value_counts()
                    st.bar_chart(quality_dist)
                
                st.divider()
                
                # Top files
                st.subheader("🏆 Top 10 Largest Extractions")
                top = df.nlargest(10, 'word_count')[['filename', 'word_count', 'page_count', 'timestamp']]
                top['timestamp'] = top['timestamp'].dt.strftime('%Y-%m-%d %H:%M')
                top.columns = ['Filename', 'Words', 'Pages', 'Date']
                st.dataframe(top, use_container_width=True, hide_index=True)
                
                st.divider()
                
                # Processing stats
                if df['processing_time'].notna().any():
                    st.subheader("⚡ Performance Metrics")
                    col_p1, col_p2, col_p3 = st.columns(3)
                    
                    avg_time = df['processing_time'].mean()
                    total_time = df['processing_time'].sum()
                    
                    col_p1.metric("Avg Time/Doc", f"{avg_time:.1f}s")
                    col_p2.metric("Total Processing Time", f"{total_time/60:.1f} min")
                    
                    if df['page_count'].sum() > 0:
                        time_per_page = total_time / df['page_count'].sum()
                        col_p3.metric("Avg Time/Page", f"{time_per_page:.1f}s")
                
                st.divider()
                
                # Recent activity
                st.subheader("🕐 Recent Activity")
                recent = df.head(10)[['filename', 'page_count', 'word_count', 'timestamp']]
                recent['timestamp'] = recent['timestamp'].dt.strftime('%Y-%m-%d %H:%M')
                recent.columns = ['Filename', 'Pages', 'Words', 'Date']
                st.dataframe(recent, use_container_width=True, hide_index=True)
                
            else:
                st.info("📭 No data available. Process some PDFs to see analytics!")
        
        except Exception as e:
            st.error(f"Error loading analytics: {e}")
    else:
        st.error("❌ Database unavailable")

# ============================================================================
# FOOTER
# ============================================================================
st.divider()
st.markdown("""
<div style="text-align: center; color: #666; padding: 1.5rem;">
    <p style="font-size: 1rem; margin-bottom: 0.5rem;">
        <strong>🚀 PDF OCR System Pro</strong> - Portfolio Project
    </p>
    <p style="font-size: 0.85rem; margin: 0;">
        Built with Python • Streamlit • Tesseract OCR • SQLite
    </p>
    <p style="font-size: 0.75rem; margin-top: 0.5rem; color: #999;">
        Developed by [Your Name] | 
        <a href="https://github.com/yourusername/project" style="color: #667eea; text-decoration: none;">GitHub</a> | 
        <a href="https://linkedin.com/in/yourprofile" style="color: #667eea; text-decoration: none;">LinkedIn</a>
    </p>
</div>
""", unsafe_allow_html=True)